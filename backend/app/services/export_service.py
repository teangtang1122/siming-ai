"""Export document generation service."""
from __future__ import annotations

import io
import json as json_module
import uuid
import zlib
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from ..core.db_helpers import get_project_or_404
from ..core.exceptions import NotFoundError, ValidationError
from ..database.models import (
    Chapter,
    Character,
    CharacterRelationship,
    OutlineNode,
    Project,
    WorldbuildingEntry,
)
from .outline_service import load_outline_nodes, outline_sort_context

EXPORT_ROOT = Path(__file__).resolve().parents[3] / "artifacts" / "exports"

def _build_chapter_text(chapter: Chapter) -> str:
    return f"{'=' * 60}\n{chapter.title}\n{'=' * 60}\n\n{chapter.content or ''}\n\n"


def _ordered_chapters(
    db: Session,
    project_id: str,
    chapter_ids: Optional[list[str]] = None,
) -> list[Chapter]:
    """Return chapters in the same outline order used by the writing workspace."""
    outline_context = outline_sort_context(load_outline_nodes(db, project_id))
    query = db.query(Chapter).filter(Chapter.project_id == project_id)
    if chapter_ids:
        unique_ids = list(dict.fromkeys(chapter_ids))
        query = query.filter(Chapter.id.in_(unique_ids))
    chapters = query.all()
    if chapter_ids and len(chapters) != len(set(chapter_ids)):
        raise ValidationError("导出章节必须属于当前作品")

    def sort_key(chapter: Chapter):
        outline_key = outline_context["sort_keys"].get(chapter.outline_node_id)
        if outline_key is None:
            return (1, (999999,), chapter.created_at or datetime.min)
        return (0, outline_key, chapter.created_at or datetime.min)

    chapters.sort(key=sort_key)
    return chapters


def _build_outline_text(nodes: list[OutlineNode]) -> str:
    lines = [f"{'=' * 60}\n大纲结构\n{'=' * 60}\n"]

    def render(node: OutlineNode, depth: int = 0):
        prefix = "  " * depth + ("- " if depth > 0 else "")
        status_map = {"pending": "[待规划]", "in_progress": "[进行中]", "completed": "[已完成]"}
        status = status_map.get(node.status, "")
        lines.append(f"{prefix}{node.title} {status}")
        if node.summary:
            lines.append(f"{'  ' * (depth + 1)}摘要: {node.summary[:200]}")
        for child in node.children:
            render(child, depth + 1)

    roots = [n for n in nodes if not n.parent_id]
    for root in roots:
        render(root)
    return "\n".join(lines) + "\n\n"


def _build_characters_text(db: Session, project_id: str) -> str:
    characters = (
        db.query(Character)
        .filter(Character.project_id == project_id)
        .order_by(Character.updated_at.desc())
        .all()
    )
    if not characters:
        return "暂无角色设定。\n\n"

    lines = [f"{'=' * 60}\n角色设定\n{'=' * 60}\n"]
    for character in characters:
        lines.append(f"\n【{character.name}】({character.role_type or '未分类'})")
        if character.appearance:
            lines.append(f"  外貌: {character.appearance}")
        if character.personality:
            lines.append(f"  性格: {character.personality}")
        if character.background:
            lines.append(f"  背景: {character.background}")
        if character.abilities:
            try:
                abilities = json_module.loads(character.abilities)
                if isinstance(abilities, list) and abilities:
                    lines.append(f"  能力: {', '.join(abilities)}")
            except (json_module.JSONDecodeError, TypeError):
                pass

    # Relationships
    relationships = (
        db.query(CharacterRelationship)
        .filter(CharacterRelationship.project_id == project_id)
        .all()
    )
    if relationships:
        lines.append(f"\n{'─' * 40}\n角色关系网络\n{'─' * 40}")
        for relationship in relationships:
            char_a = db.query(Character).filter(Character.id == relationship.character_a_id).first()
            char_b = db.query(Character).filter(Character.id == relationship.character_b_id).first()
            name_a = char_a.name if char_a else relationship.character_a_id[:8]
            name_b = char_b.name if char_b else relationship.character_b_id[:8]
            lines.append(f"  {name_a} → {name_b}: {relationship.relationship_type}")
            if relationship.description:
                lines.append(f"    说明: {relationship.description}")

    return "\n".join(lines) + "\n\n"


def _build_worldbuilding_text(db: Session, project_id: str) -> str:
    entries = (
        db.query(WorldbuildingEntry)
        .filter(WorldbuildingEntry.project_id == project_id)
        .order_by(WorldbuildingEntry.dimension.asc(), WorldbuildingEntry.sort_order.asc())
        .all()
    )
    if not entries:
        return "暂无世界观设定。\n\n"

    dim_labels = {
        "geography": "地理", "history": "历史", "factions": "势力",
        "power_system": "规则体系", "races": "种族", "culture": "文化",
    }
    lines = [f"{'=' * 60}\n世界观设定\n{'=' * 60}\n"]
    current_dim = None
    for entry in entries:
        if entry.dimension != current_dim:
            current_dim = entry.dimension
            lines.append(f"\n## {dim_labels.get(current_dim, current_dim)}")
        lines.append(f"\n### {entry.title}")
        lines.append(entry.content)
    return "\n".join(lines) + "\n\n"


def _generate_export_content(
    db: Session,
    project_id: str,
    scope: str,
    chapter_ids: Optional[list[str]] = None,
    include_outline: bool = False,
    include_characters: bool = False,
    include_worldbuilding: bool = False,
) -> tuple[str, str]:
    """Generate export content. Returns (filename, content)."""
    project = get_project_or_404(db, project_id)
    safe_title = project.title.replace("/", "_").replace("\\", "_")[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.txt"
    selected_chapter_ids = chapter_ids or []
    include_chapters = scope in ("chapters", "all", "single", "selected") or bool(selected_chapter_ids)
    if scope in ("single", "selected") and not selected_chapter_ids:
        raise ValidationError("单章或指定章节导出必须提供 chapter_ids")

    sections = [f"{project.title}\n{'=' * 60}\n"]
    if project.description:
        sections.append(f"简介: {project.description}\n")

    if include_chapters:
        sections.append(f"\n{'=' * 60}\n正文\n{'=' * 60}\n")
        chapters = _ordered_chapters(db, project_id, selected_chapter_ids)
        if chapters:
            for chapter in chapters:
                sections.append(_build_chapter_text(chapter))
        else:
            sections.append("暂无章节内容。\n")

    if scope in ("outline", "all") or include_outline:
        nodes = (
            db.query(OutlineNode)
            .filter(OutlineNode.project_id == project_id)
            .order_by(OutlineNode.sort_order.asc(), OutlineNode.created_at.asc())
            .all()
        )
        sections.append(_build_outline_text(nodes))

    if scope in ("characters", "all") or include_characters:
        sections.append(_build_characters_text(db, project_id))

    if scope in ("worldbuilding", "all") or include_worldbuilding:
        sections.append(_build_worldbuilding_text(db, project_id))

    content = "\n".join(sections)
    return filename, content


def _generate_docx(
    db: Session,
    project_id: str,
    scope: str,
    chapter_ids: Optional[list[str]] = None,
    include_outline: bool = False,
    include_characters: bool = False,
    include_worldbuilding: bool = False,
) -> tuple[str, io.BytesIO]:
    """Generate a Word (.docx) document. Returns (filename, BytesIO buffer)."""
    project = get_project_or_404(db, project_id)
    safe_title = project.title.replace("/", "_").replace("\\", "_")[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.docx"
    selected_chapter_ids = chapter_ids or []
    include_chapters = scope in ("chapters", "all", "single", "selected") or bool(selected_chapter_ids)
    if scope in ("single", "selected") and not selected_chapter_ids:
        raise ValidationError("单章或指定章节导出必须提供 chapter_ids")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    title_para = doc.add_heading(project.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project.description:
        desc = doc.add_paragraph(project.description)
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc.runs[0].font.size = Pt(12)

    doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    if include_chapters:
        doc.add_heading("正文", level=1)
        chapters = _ordered_chapters(db, project_id, selected_chapter_ids)
        if chapters:
            for ch in chapters:
                doc.add_heading(ch.title, level=2)
                if ch.content:
                    for paragraph_text in ch.content.split("\n"):
                        if paragraph_text.strip():
                            doc.add_paragraph(paragraph_text.strip())
                doc.add_paragraph()
        else:
            doc.add_paragraph("暂无章节内容。")

    if scope in ("outline", "all") or include_outline:
        doc.add_page_break()
        doc.add_heading("大纲结构", level=1)
        nodes = (
            db.query(OutlineNode)
            .filter(OutlineNode.project_id == project_id)
            .order_by(OutlineNode.sort_order.asc(), OutlineNode.created_at.asc())
            .all()
        )

        def render_outline(node: OutlineNode, depth: int = 0):
            level = min(depth + 2, 5)
            status_map = {"pending": "[待规划]", "in_progress": "[进行中]", "completed": "[已完成]"}
            heading = doc.add_heading(f"{node.title} {status_map.get(node.status, '')}", level=level)
            if node.summary:
                doc.add_paragraph(f"摘要: {node.summary[:200]}")
            for child in node.children:
                render_outline(child, depth + 1)

        roots = [n for n in nodes if not n.parent_id]
        for root in roots:
            render_outline(root)

    if scope in ("characters", "all") or include_characters:
        doc.add_page_break()
        doc.add_heading("角色设定", level=1)
        characters = (
            db.query(Character)
            .filter(Character.project_id == project_id)
            .order_by(Character.updated_at.desc())
            .all()
        )
        if characters:
            for character in characters:
                doc.add_heading(f"{character.name} ({character.role_type or '未分类'})", level=2)
                if character.appearance:
                    doc.add_paragraph("外貌", style="List Bullet")
                    doc.add_paragraph(character.appearance)
                if character.personality:
                    doc.add_paragraph("性格", style="List Bullet")
                    doc.add_paragraph(character.personality)
                if character.background:
                    doc.add_paragraph("背景", style="List Bullet")
                    doc.add_paragraph(character.background)
                if character.abilities:
                    try:
                        abilities = json_module.loads(character.abilities)
                        if isinstance(abilities, list) and abilities:
                            doc.add_paragraph("能力", style="List Bullet")
                            doc.add_paragraph(", ".join(abilities))
                    except (json_module.JSONDecodeError, TypeError):
                        pass

            relationships = (
                db.query(CharacterRelationship)
                .filter(CharacterRelationship.project_id == project_id)
                .all()
            )
            if relationships:
                doc.add_heading("角色关系网络", level=2)
                for rel in relationships:
                    char_a = db.query(Character).filter(Character.id == rel.character_a_id).first()
                    char_b = db.query(Character).filter(Character.id == rel.character_b_id).first()
                    name_a = char_a.name if char_a else rel.character_a_id[:8]
                    name_b = char_b.name if char_b else rel.character_b_id[:8]
                    text = f"{name_a} → {name_b}: {rel.relationship_type}"
                    if rel.description:
                        text += f" ({rel.description})"
                    doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph("暂无角色设定。")

    if scope in ("worldbuilding", "all") or include_worldbuilding:
        doc.add_page_break()
        doc.add_heading("世界观设定", level=1)
        entries = (
            db.query(WorldbuildingEntry)
            .filter(WorldbuildingEntry.project_id == project_id)
            .order_by(WorldbuildingEntry.dimension.asc(), WorldbuildingEntry.sort_order.asc())
            .all()
        )
        if entries:
            dim_labels = {
                "geography": "地理", "history": "历史", "factions": "势力",
                "power_system": "规则体系", "races": "种族", "culture": "文化",
            }
            current_dim = None
            for entry in entries:
                if entry.dimension != current_dim:
                    current_dim = entry.dimension
                    doc.add_heading(dim_labels.get(current_dim, current_dim), level=2)
                doc.add_heading(entry.title, level=3)
                doc.add_paragraph(entry.content)
        else:
            doc.add_paragraph("暂无世界观设定。")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return filename, buf


def _generate_pdf(
    db: Session,
    project_id: str,
    scope: str,
    chapter_ids: Optional[list[str]] = None,
    include_outline: bool = False,
    include_characters: bool = False,
    include_worldbuilding: bool = False,
) -> tuple[str, io.BytesIO]:
    """Generate a PDF document with Chinese text (CID font, UTF-16BE). Returns (filename, BytesIO buffer)."""
    project = get_project_or_404(db, project_id)
    safe_title = project.title.replace("/", "_").replace("\\", "_")[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.pdf"

    _, txt_content = _generate_export_content(
        db,
        project_id,
        scope,
        chapter_ids=chapter_ids,
        include_outline=include_outline,
        include_characters=include_characters,
        include_worldbuilding=include_worldbuilding,
    )

    objects: list[bytes] = []
    offsets: list[int] = []

    def add_obj(data: bytes) -> int:
        obj_num = len(objects) + 1
        offsets.append(0)
        objects.append(f"{obj_num} 0 obj\n".encode() + data + b"\nendobj\n")
        return obj_num

    page_w, page_h = 595.28, 841.89  # A4 in points
    margin_l, margin_r, margin_t, margin_b = 50, 50, 50, 50
    font_size = 11
    line_height = 17.6
    y_start = page_h - margin_t
    max_line_w = page_w - margin_l - margin_r

    # Build CID font with Identity-H encoding for Unicode
    cid_font_obj = add_obj(
        b"<</Type /Font /Subtype /Type0 /BaseFont /MS-Gothic /Encoding /Identity-H "
        b"/DescendantFonts [<</Type /Font /Subtype /CIDFontType2 "
        b"/BaseFont /MS-Gothic /CIDSystemInfo <</Registry (Adobe) /Ordering (Identity) /Supplement 0>> "
        b"/DW 1000 /W [0 [600]]>>]>>"
    )

    pages_content: list[tuple[bytes, bytes]] = []  # (content_stream, resources_dict)
    current_lines: list[str] = []
    current_y = y_start

    all_lines = txt_content.split("\n")

    def estimate_line_width(text: str) -> float:
        """Rough width estimate: CJK chars ~2x Latin chars."""
        w = 0.0
        for ch in text:
            if '一' <= ch <= '鿿' or '　' <= ch <= '〿' or '＀' <= ch <= '￯':
                w += font_size * 1.0
            else:
                w += font_size * 0.5
        return w

    def wrap_line(text: str, max_w: float) -> list[str]:
        if not text:
            return [""]
        result: list[str] = []
        current = ""
        current_w = 0.0
        for ch in text:
            ch_w = font_size * (1.0 if ord(ch) > 127 else 0.5)
            if current_w + ch_w > max_w and current:
                result.append(current)
                current = ""
                current_w = 0.0
            current += ch
            current_w += ch_w
        if current:
            result.append(current)
        return result or [""]

    for line in all_lines:
        wrapped = wrap_line(line, max_line_w)
        for segment in wrapped:
            if current_y < margin_b + line_height:
                stream = _build_pdf_stream(current_lines, margin_l, y_start, font_size, line_height, cid_font_obj)
                pages_content.append((stream, cid_font_obj))
                current_lines = []
                current_y = y_start
            current_lines.append(segment)
            current_y -= line_height

    if current_lines:
        stream = _build_pdf_stream(current_lines, margin_l, y_start, font_size, line_height, cid_font_obj)
        pages_content.append((stream, cid_font_obj))

    page_obj_ids = []
    for i, (stream_bytes, _) in enumerate(pages_content):
        compressed = zlib.compress(stream_bytes)
        content_id = add_obj(
            f"<</Length {len(compressed)} /Filter /FlateDecode>>\nstream\n".encode()
            + compressed + b"\nendstream"
        )
        resources = f"<</Font <</CIDFont {cid_font_obj} 0 R>>>>".encode()
        page_id = add_obj(
            b"<</Type /Page /Parent 0 0 R /MediaBox [0 0 %f %f] /Contents %d 0 R /Resources %s>>" % (
                page_w, page_h, content_id, resources
            )
        )
        page_obj_ids.append(page_id)

    kids_refs = b" ".join(f"{pid} 0 R".encode() for pid in page_obj_ids)
    pages_id = add_obj(
        b"<</Type /Pages /Kids [" + kids_refs + b"] /Count " + str(len(page_obj_ids)).encode() + b">>"
    )

    # Patch parent references in page objects
    for pid in page_obj_ids:
        for i, obj_data in enumerate(objects):
            if obj_data.startswith(f"{pid} 0 obj".encode()):
                objects[i] = obj_data.replace(b"/Parent 0 0 R", f"/Parent {pages_id} 0 R".encode())
                break

    catalog_id = add_obj(f"<</Type /Catalog /Pages {pages_id} 0 R>>".encode())

    buf = io.BytesIO()
    buf.write(b"%PDF-1.4\n")
    buf.write(b"%\x80\x80\x80\x80\n")

    current_offset = buf.tell()
    for i, obj_data in enumerate(objects):
        offsets[i] = current_offset
        buf.write(obj_data)
        current_offset = buf.tell()

    xref_offset = current_offset
    buf.write(b"xref\n")
    buf.write(f"0 {len(objects) + 1}\n".encode())
    buf.write(b"0000000000 65535 f \n")
    for off in offsets:
        buf.write(f"{off:010d} 00000 n \n".encode())

    buf.write(b"trailer\n")
    buf.write(b"<</Size %d /Root %d 0 R>>\n" % (len(objects) + 1, catalog_id))
    buf.write(b"startxref\n")
    buf.write(f"{xref_offset}\n".encode())
    buf.write(b"%%EOF")

    buf.seek(0)
    return filename, buf


def _build_pdf_stream(lines: list[str], margin: float, y_start: float, font_size: int, line_height: float, font_obj_id: int) -> bytes:
    """Build a PDF page content stream with UTF-16BE hex-encoded Chinese text."""
    parts = [b"BT\n"]
    parts.append(f"/CIDFont {font_size} Tf\n".encode())
    y = y_start
    for line in lines:
        line = line.replace("\r", "")
        hex_str = line.encode("utf-16-be").hex()
        parts.append(f"1 0 0 1 {margin} {y} Tm\n".encode())
        parts.append(f"<{hex_str}> Tj\n".encode())
        y -= line_height
    parts.append(b"ET\n")
    return b"".join(parts)


def _store_export_file(project_id: str, filename: str, data: bytes, media_type: str, export_format: str) -> dict:
    file_id = str(uuid.uuid4())
    project_dir = EXPORT_ROOT / project_id
    project_dir.mkdir(parents=True, exist_ok=True)

    suffix = Path(filename).suffix or f".{export_format}"
    stored_filename = f"{file_id}{suffix}"
    (project_dir / stored_filename).write_bytes(data)

    metadata = {
        "file_id": file_id,
        "project_id": project_id,
        "filename": filename,
        "stored_filename": stored_filename,
        "format": export_format,
        "media_type": media_type,
        "size": len(data),
        "created_at": datetime.now().isoformat(),
        "download_url": f"/api/v1/projects/{project_id}/export/download/{file_id}",
    }
    (project_dir / f"{file_id}.json").write_text(
        json_module.dumps(metadata, ensure_ascii=False),
        encoding="utf-8",
    )
    return metadata


def _load_export_metadata(project_id: str, file_id: str) -> dict:
    try:
        uuid.UUID(file_id)
    except ValueError:
        raise NotFoundError("导出文件不存在")

    metadata_path = EXPORT_ROOT / project_id / f"{file_id}.json"
    if not metadata_path.exists():
        raise NotFoundError("导出文件不存在")
    return json_module.loads(metadata_path.read_text(encoding="utf-8"))

