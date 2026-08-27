"""Regression tests for imported chapter persistence."""

import base64
import contextlib
import json
import os
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock, patch

os.environ["DATABASE_URL"] = "sqlite:///./test_novel_agent.db"

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.core.utils import count_words
from app.database.models import Chapter, Project
from app.database.session import Base, SessionLocal, engine
from app.main import app
from app.services.import_service import _parse_raw_file

API_PREFIX = "/api/v1"


class ImporterTestCase(unittest.TestCase):
    """Import confirmation should create valid project-owned chapters."""

    @classmethod
    def setUpClass(cls):
        Base.metadata.create_all(bind=engine)
        cls.client = TestClient(app)

    @classmethod
    def tearDownClass(cls):
        Base.metadata.drop_all(bind=engine)
        with contextlib.suppress(OSError):
            os.remove("test_novel_agent.db")

    def create_project(self, title: str) -> str:
        response = self.client.post(f"{API_PREFIX}/projects", json={"title": title})
        self.assertEqual(response.status_code, 200)
        return response.json()["data"]["id"]

    def create_outline_node(self, project_id: str, title: str) -> str:
        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/outline",
            json={"title": title, "node_type": "chapter"},
        )
        self.assertEqual(response.status_code, 200)
        return response.json()["data"]["id"]

    def test_confirm_import_returns_ids_and_persists_word_counts(self):
        project_id = self.create_project("Import Project")
        outline_id = self.create_outline_node(project_id, "Import Target")
        first = "one two"
        second = "three four five"
        text = f"{first}\n{second}"

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/confirm",
            json={
                "text": text,
                "outline_node_id": outline_id,
                "splits": [
                    {"title": "One", "start_char": 0, "end_char": len(first), "preview": first},
                    {
                        "title": "Two",
                        "start_char": len(first) + 1,
                        "end_char": len(text),
                        "preview": second,
                    },
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        chapters = response.json()["data"]["chapters"]
        expected_counts = [count_words(first), count_words(second)]
        self.assertEqual([c["word_count"] for c in chapters], expected_counts)
        self.assertTrue(all(c["id"] for c in chapters))

        db = SessionLocal()
        try:
            stored = (
                db.query(Chapter)
                .filter(Chapter.project_id == project_id)
                .order_by(Chapter.sort_order.asc(), Chapter.created_at.asc(), Chapter.id.asc())
                .all()
            )
            self.assertEqual(len(stored), 2)
            self.assertEqual([chapter.word_count for chapter in stored], expected_counts)
            self.assertEqual([chapter.sort_order for chapter in stored], [1000, 2000])
            self.assertTrue(all(chapter.outline_node_id == outline_id for chapter in stored))
        finally:
            db.close()

    def test_confirm_import_rejects_cross_project_outline_node(self):
        project_a = self.create_project("Project A")
        project_b = self.create_project("Project B")
        foreign_outline_id = self.create_outline_node(project_b, "Foreign Target")

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_a}/import/confirm",
            json={
                "text": "one two",
                "outline_node_id": foreign_outline_id,
                "splits": [],
            },
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("当前作品", response.json()["message"])

    def test_import_file_alias_parses_txt(self):
        project_id = self.create_project("Upload Project")

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/file",
            files={"file": ("sample.txt", "第一章\n正文".encode(), "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["filename"], "sample.txt")
        self.assertEqual(data["format"], "txt")
        self.assertIn("第一章", data["text"])

    def test_import_file_alias_parses_markdown_without_losing_markup(self):
        project_id = self.create_project("Markdown Upload Project")
        text = "# 第一章 风起\n\n这里有 **加粗正文**。"

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/file",
            files={"file": ("sample.md", text.encode("utf-8"), "text/markdown")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["format"], "md")
        self.assertEqual(data["text"], text)
        self.assertEqual(data["encoding"], "UTF-8")

    def test_import_file_detects_gb18030(self):
        project_id = self.create_project("GB18030 Project")
        text = "第一章 风起\n陆糖看见归墟阵重新亮起。"

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/file",
            files={"file": ("gb.txt", text.encode("gb18030"), "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["encoding"], "GB18030")
        self.assertEqual(data["text"], text)
        self.assertNotIn("\ufffd", data["text"])

    def test_import_file_detects_utf16le_without_bom(self):
        project_id = self.create_project("UTF16 Project")
        text = "第一章 风起\n这是 UTF-16LE 正文。"

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/file",
            files={"file": ("utf16.txt", text.encode("utf-16-le"), "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["encoding"], "UTF-16LE")
        self.assertEqual(data["text"], text)

    def test_shared_android_encoding_fixtures_match_pc_decoder(self):
        fixture_path = (
            Path(__file__).resolve().parents[2]
            / "contracts"
            / "novel-import-encoding-fixtures.json"
        )
        fixture = json.loads(fixture_path.read_text(encoding="utf-8"))

        for case in fixture["cases"]:
            with self.subTest(case=case["name"]):
                data = _parse_raw_file(case["filename"], base64.b64decode(case["base64"]))
                self.assertEqual(data["encoding"], case["expected_encoding"])
                self.assertEqual(data["text"], case["text"])
                self.assertNotIn("\ufffd", data["text"])

    def test_atomic_project_file_import_creates_project_and_all_chapters(self):
        text = (
            "第一章 风起\n"
            + "这里是第一章正文。" * 10
            + "\n\n第二章 云涌\n"
            + "这里是第二章正文。" * 10
        )

        response = self.client.post(
            f"{API_PREFIX}/import/project-file",
            files={"file": ("批量导入.txt", text.encode("gb18030"), "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["encoding"], "GB18030")
        self.assertEqual(data["total"], 2)
        project_id = data["project_id"]

        db = SessionLocal()
        try:
            project = db.query(Project).filter(Project.id == project_id).one_or_none()
            chapters = (
                db.query(Chapter)
                .filter(Chapter.project_id == project_id)
                .order_by(Chapter.sort_order.asc())
                .all()
            )
            self.assertIsNotNone(project)
            self.assertEqual(project.title, "批量导入")
            self.assertEqual(len(chapters), 2)
            self.assertEqual([chapter.sort_order for chapter in chapters], [1000, 2000])
        finally:
            db.close()

    def test_atomic_project_file_import_uses_markdown_headings_as_boundaries(self):
        text = (
            "# 第一章 风起\n\n"
            + "这里是 **第一章** 正文。" * 10
            + "\n\n## 第二章 云涌\n\n"
            + "这里是 _第二章_ 正文。" * 10
        )

        response = self.client.post(
            f"{API_PREFIX}/import/project-file",
            files={"file": ("Markdown小说.md", text.encode("utf-8"), "text/markdown")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["format"], "md")
        self.assertEqual(data["total"], 2)
        self.assertEqual(
            [chapter["title"] for chapter in data["chapters"]],
            ["第一章 风起", "第二章 云涌"],
        )
        db = SessionLocal()
        try:
            chapters = (
                db.query(Chapter)
                .filter(Chapter.project_id == data["project_id"])
                .order_by(Chapter.sort_order.asc())
                .all()
            )
            self.assertEqual(len(chapters), 2)
            self.assertTrue(chapters[0].content.startswith("# 第一章"))
            self.assertIn("**第一章**", chapters[0].content)
        finally:
            db.close()

    def test_atomic_project_file_import_accepts_android_docx_payload(self):
        document = DocxDocument()
        document.add_paragraph("第一章 风起")
        document.add_paragraph("这里是第一章正文。" * 10)
        document.add_paragraph("第二章 云涌")
        document.add_paragraph("这里是第二章正文。" * 10)
        buffer = BytesIO()
        document.save(buffer)

        data = _parse_raw_file("手机导入.docx", buffer.getvalue())
        self.assertEqual(data["encoding"], "DOCX")
        self.assertEqual(data["format"], "docx")
        self.assertIn("第一章 风起", data["text"])
        self.assertIn("第二章 云涌", data["text"])

    def test_import_preview_ignores_sentence_like_chapter_prefixes_in_body(self):
        project_id = self.create_project("Body Prefix Project")
        text = (
            "第一章 风起！\n"
            + "第一章正文。这里仍然属于正文，不是新章节。" * 6
            + "\n\n第二章 云涌\n"
            + "第二章正文继续。" * 6
        )

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/preview",
            json={"text": text},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["total"], 2)
        self.assertEqual(
            [item["title"] for item in data["splits"]],
            ["第一章 风起！", "第二章 云涌"],
        )

    def test_import_preview_uses_regex_chapter_boundaries_without_llm(self):
        project_id = self.create_project("Preview Project")
        text = (
            "第一章 风起\n"
            + "这里是第一章正文。" * 10
            + "\n\n第二章 云涌\n"
            + "这里是第二章正文。" * 10
        )

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/preview",
            json={"text": text},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["method"], "regex")
        self.assertEqual(data["total"], 2)
        self.assertEqual([item["title"] for item in data["splits"]], ["第一章 风起", "第二章 云涌"])

    @patch("app.routers.importer.LLMGateway.chat_completion", new_callable=AsyncMock)
    def test_import_preview_uses_chunked_llm_corrections(self, mock_chat):
        project_id = self.create_project("LLM Preview Project")
        text = (
            "第一章 风起\n"
            + "这里是第一章正文。" * 10
            + "\n\n第二章 云涌\n"
            + "这里是第二章正文。" * 10
        )
        mock_chat.return_value = {
            "content": json.dumps(
                [
                    {
                        "title": "第一章 风起（校正）",
                        "start_char": 0,
                        "end_char": text.index("第二章 云涌"),
                        "preview": "这里是第一章正文。",
                    },
                    {
                        "title": "第二章 云涌（校正）",
                        "start_char": text.index("第二章 云涌"),
                        "end_char": len(text),
                        "preview": "这里是第二章正文。",
                    },
                ],
                ensure_ascii=False,
            )
        }

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/preview",
            json={"text": text, "model": "openai:test"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["method"], "regex+chunked-llm")
        self.assertEqual(data["failed_blocks"], 0)
        self.assertEqual(
            [item["title"] for item in data["splits"]],
            ["第一章 风起（校正）", "第二章 云涌（校正）"],
        )
        self.assertTrue(all(item["source"] == "llm" for item in data["splits"]))
        self.assertEqual(mock_chat.await_count, 1)

    @patch("app.routers.importer.asyncio.sleep", new_callable=AsyncMock)
    @patch("app.routers.importer.LLMGateway.chat_completion", new_callable=AsyncMock)
    def test_import_preview_marks_failed_llm_blocks_for_manual_review(self, mock_chat, mock_sleep):
        project_id = self.create_project("Failed LLM Preview Project")
        text = "\n\n".join(f"第{i}章 标题{i}\n" + f"这里是第{i}章正文。" * 8 for i in range(1, 5))
        mock_chat.side_effect = RuntimeError("llm unavailable")

        response = self.client.post(
            f"{API_PREFIX}/projects/{project_id}/import/preview",
            json={"text": text, "model": "openai:test"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["method"], "regex+chunked-llm")
        self.assertEqual(data["failed_blocks"], 2)
        self.assertTrue(data["needs_review"])
        self.assertEqual(data["total"], 4)
        self.assertTrue(all(item["needs_review"] for item in data["splits"]))
        self.assertEqual(mock_chat.await_count, 6)
        self.assertEqual(mock_sleep.await_count, 4)
